"""
Markdown to Word Converter
==========================
Converts markdown files (.md) to Word documents (.docx).
Works for CVs, Cover Letters, and general markdown documents.

USAGE:
    python 01_md_to_docx.py <input.md> <output.docx>
    
    Or run without arguments to use the interactive prompt.

REQUIREMENTS:
    pip install python-docx

AUTHOR: Job Search Tools Collection
"""

import re
import sys
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("ERROR: python-docx not installed.")
    print("Please run: pip install python-docx")
    sys.exit(1)


def add_formatted_text(paragraph, text):
    """
    Parse markdown bold (**text**) and links [text](url) and add to paragraph.
    """
    # Split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle links [text](url) - extract just the text
            part = re.sub(r'\[(.*?)\]\((.*?)\)', r'\1', part)
            if part:
                paragraph.add_run(part)


def convert_md_to_docx(md_file, docx_file):
    """
    Convert a markdown file to a Word document.
    Handles headings, bullet points, bold text, and horizontal rules.
    """
    print(f"Converting: {md_file}")
    print(f"Output:     {docx_file}")
    
    # Read markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Create document with basic styling
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for line in lines:
        line = line.strip()
        
        # Skip empty lines (but could add spacing if desired)
        if not line:
            continue
        
        # Title (# Heading)
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
        # Heading 2 (## Heading)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
            
        # Heading 3 (### Heading)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
            
        # Bullet points (* item or - item)
        elif line.startswith('* ') or line.startswith('- ') or \
             line.startswith('*   ') or line.startswith('-   '):
            clean_line = line.lstrip('*- ').strip()
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, clean_line)
            
        # Horizontal rule (---)
        elif line.startswith('---'):
            # Skip horizontal rules in Word (they don't translate well)
            pass
            
        # Normal text
        else:
            # Check if it's a contact info line (contains |) - center it
            if '|' in line:
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                add_formatted_text(p, line)
            else:
                p = doc.add_paragraph()
                add_formatted_text(p, line)

    # Save document
    doc.save(docx_file)
    print(f"✓ Successfully created: {docx_file}")


def main():
    """Main entry point with argument handling."""
    if len(sys.argv) >= 3:
        # Arguments provided
        md_file = sys.argv[1]
        docx_file = sys.argv[2]
    elif len(sys.argv) == 2:
        # Only input provided, auto-generate output name
        md_file = sys.argv[1]
        docx_file = Path(md_file).with_suffix('.docx')
    else:
        # Interactive mode
        print("=" * 50)
        print("  Markdown to Word Converter")
        print("=" * 50)
        md_file = input("Enter markdown file path: ").strip().strip('"')
        
        # Auto-suggest output filename
        suggested = Path(md_file).with_suffix('.docx')
        docx_file = input(f"Output path [{suggested}]: ").strip().strip('"')
        if not docx_file:
            docx_file = suggested
    
    # Validate input file exists
    if not os.path.exists(md_file):
        print(f"ERROR: File not found: {md_file}")
        sys.exit(1)
    
    convert_md_to_docx(md_file, docx_file)


if __name__ == "__main__":
    main()
